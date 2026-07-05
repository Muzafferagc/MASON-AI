"""
documents.py - Belge hafizasi (dosya yukleme + RAG)

Mason'a dosya yukleyebilirsin (ChatGPT/Gemini gibi): PDF, Word, Excel/CSV,
metin, kod, GORSEL ve SES dosyalari. Her dosya:
  1. Icerigi cikartilir (metne cevrilir)
       - PDF/docx/xlsx: ilgili kutuphane ile
       - txt/md/kod/csv: dogrudan okunur
       - gorsel: Gemini "goru" (vision) ile okunur/aciklanir
       - ses: faster-whisper ile yaziya cevrilir (Faz 2 altyapisi)
  2. Anlamli PARCALARA (chunk) bolunur
  3. Her parcaya embedding cikarilir ve veritabanina kaydedilir

Sonra bir soru sordugunda, soruyla ANLAMCA en ilgili parcalar bulunup
Mason'un promptuna eklenir (RAG). Boylece 100 sayfalik bir PDF'in
tamamini degil, sadece ilgili 5-6 parcasini okur — hem hizli hem ucretsiz
kota dostu. Belgeler kalicidir; sohbet kapansa da hafizada kalir.
"""
import base64
import json
import os
from pathlib import Path

from .database import get_conn, rows_to_dicts
from .embeddings import embed_text, cosine_similarity

# ---- Desteklenen dosya turleri (uzanti -> tur etiketi) ----
TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".rst", ".log",
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".c", ".cpp", ".h", ".hpp",
    ".cs", ".go", ".rs", ".rb", ".php", ".swift", ".kt", ".sql", ".sh", ".bat",
    ".ps1", ".html", ".htm", ".css", ".json", ".yaml", ".yml", ".xml", ".ini",
    ".cfg", ".toml", ".env", ".csv", ".tsv",
}
PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx", ".xlsm", ".xls"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff"}
AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".aac", ".opus"}

SUPPORTED_EXTS = (
    TEXT_EXTS | PDF_EXTS | DOCX_EXTS | XLSX_EXTS | IMAGE_EXTS | AUDIO_EXTS
)

# Parcalama ayarlari
CHUNK_SIZE = 1200        # bir parcadaki yaklasik karakter sayisi
CHUNK_OVERLAP = 150      # parcalarin birbirine binmesi (baglam kopmasin)
# Retrieval: her soruda prompt'a konacak en fazla parca sayisi ve toplam karakter
TOP_CHUNKS = 8
MAX_CONTEXT_CHARS = 9000
# Belge sayisi azken embedding aramasi yapmadan en yeni parcalar kullanilir
SEMANTIC_THRESHOLD = 12


def filetype_of(path: str) -> str:
    """Uzantiya gore kisa tur etiketi (arayuzde ikon/rozet icin)."""
    ext = Path(path).suffix.lower()
    if ext in PDF_EXTS:
        return "pdf"
    if ext in DOCX_EXTS:
        return "word"
    if ext in XLSX_EXTS:
        return "excel"
    if ext in IMAGE_EXTS:
        return "image"
    if ext in AUDIO_EXTS:
        return "audio"
    if ext in TEXT_EXTS:
        return "text"
    return "other"


def is_supported(path: str) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_EXTS


# ---------- Icerik cikartma ----------

def extract_text(path: str, config: dict | None = None) -> str:
    """Dosyanin icerigini metne cevirir. Desteklenmeyen/okunamayan
    dosyada aciklayici bir hata metni dondurur (uygulama cokmez)."""
    ext = Path(path).suffix.lower()
    try:
        if ext in TEXT_EXTS:
            return _read_text_file(path)
        if ext in PDF_EXTS:
            return _extract_pdf(path)
        if ext in DOCX_EXTS:
            return _extract_docx(path)
        if ext in XLSX_EXTS:
            return _extract_xlsx(path)
        if ext in IMAGE_EXTS:
            return _extract_image(path, config or {})
        if ext in AUDIO_EXTS:
            return _extract_audio(path, config or {})
    except MissingDependency as e:
        return f"[Bu dosya turu icin gerekli kutuphane kurulu degil: {e}]"
    except Exception as e:  # noqa: BLE001 - hicbir dosya uygulamayi cokertmesin
        return f"[Dosya okunamadi: {e}]"
    return "[Desteklenmeyen dosya turu]"


class MissingDependency(Exception):
    """Opsiyonel bir cikartma kutuphanesi kurulu degil."""


def _read_text_file(path: str) -> str:
    """Duz metin/kod/csv dosyalarini kodlama tahmini ile okur."""
    data = Path(path).read_bytes()
    for enc in ("utf-8", "utf-8-sig", "cp1254", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # eski isim
        except ImportError as e:
            raise MissingDependency("pypdf (pip install pypdf)") from e
    reader = PdfReader(path)
    parts = []
    for i, page in enumerate(reader.pages, 1):
        txt = page.extract_text() or ""
        if txt.strip():
            parts.append(f"[sayfa {i}]\n{txt}")
    text = "\n\n".join(parts).strip()
    return text or "[PDF metin icermiyor (taranmis/goruntu olabilir)]"


def _extract_docx(path: str) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise MissingDependency("python-docx (pip install python-docx)") from e
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    # Tablolari da al
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip() or "[Word belgesi bos]"


def _extract_xlsx(path: str) -> str:
    try:
        import openpyxl
    except ImportError as e:
        raise MissingDependency("openpyxl (pip install openpyxl)") from e
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[sayfa: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts).strip() or "[Excel dosyasi bos]"


def _extract_image(path: str, config: dict) -> str:
    """Gorseli Gemini 'goru' (vision) ile okur: icindeki metni ve gorsel
    icerigi Turkce olarak aciklar. Gemini anahtari yoksa uyari dondurur."""
    api_key = config.get("gemini_api_key")
    provider = config.get("provider", "gemini")
    if not api_key or provider == "ollama":
        return ("[Gorsel icerigi okunamadi. Gorsellerden metin/aciklama cikarmak "
                "icin Ayarlar'dan bir Gemini API anahtari girmelisin (ucretsiz).]")
    import requests
    mime = _image_mime(path)
    b64 = base64.b64encode(Path(path).read_bytes()).decode("ascii")
    model = config.get("gemini_model", "gemini-2.5-flash")
    url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
           f"{model}:generateContent?key={api_key}")
    prompt = ("Bu gorseli detayli anlat. Icinde yazi/metin varsa AYNEN yaz "
              "(OCR). Grafik/tablo varsa verilerini aktar. Turkce yanitla.")
    body = {"contents": [{"parts": [
        {"text": prompt},
        {"inline_data": {"mime_type": mime, "data": b64}},
    ]}]}
    resp = requests.post(url, json=body, timeout=120)
    if resp.status_code != 200:
        return f"[Gorsel okunamadi (Gemini HTTP {resp.status_code})]"
    try:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    except (KeyError, IndexError):
        return "[Gorsel okundu ama bos yanit dondu]"


def _image_mime(path: str) -> str:
    ext = Path(path).suffix.lower()
    return {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".webp": "image/webp", ".gif": "image/gif", ".bmp": "image/bmp",
        ".tiff": "image/tiff",
    }.get(ext, "image/png")


def _extract_audio(path: str, config: dict) -> str:
    """Ses dosyasini yaziya cevirir (faster-whisper, yerel & ucretsiz)."""
    try:
        from faster_whisper import WhisperModel
    except ImportError as e:
        raise MissingDependency(
            "faster-whisper (pip install faster-whisper)") from e
    model_size = config.get("whisper_model", "small")
    lang = config.get("stt_language", "tr")
    model = WhisperModel(model_size, device="cpu", compute_type="int8")
    segments, _info = model.transcribe(
        path, language=None if lang == "auto" else lang)
    text = " ".join(seg.text.strip() for seg in segments).strip()
    return text or "[Seste konusma algilanmadi]"


# ---------- Parcalama ----------

def chunk_text(text: str, size: int = CHUNK_SIZE,
               overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Metni, mumkun oldugunca paragraf/satir sinirlarindan bolerek
    ~'size' karakterlik ust uste binen parcalara ayirir."""
    text = (text or "").strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]
    chunks, start, n = [], 0, len(text)
    while start < n:
        end = min(start + size, n)
        if end < n:
            # En yakin paragraf/cumle/bosluk sinirinda kes (kelimeyi bolme)
            window = text[start:end]
            for sep in ("\n\n", "\n", ". ", " "):
                idx = window.rfind(sep)
                if idx > size * 0.5:
                    end = start + idx + len(sep)
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        start = max(end - overlap, start + 1)
    return chunks


# ---------- Kaydetme (ingest) ----------

def ingest(path: str, config: dict | None = None,
           stored_path: str | None = None) -> dict:
    """Bir dosyayi okur, parcalara boler, embedding cikarir ve veritabanina
    kaydeder. Ozet bilgi dondurur: {ok, id, filename, chunks, chars, ...}."""
    p = Path(path)
    if not p.exists():
        return {"ok": False, "filename": p.name, "error": "Dosya bulunamadi"}
    if not is_supported(str(p)):
        return {"ok": False, "filename": p.name,
                "error": f"Desteklenmeyen tur: {p.suffix}"}

    text = extract_text(str(p), config)
    chunks = chunk_text(text)
    if not chunks:
        return {"ok": False, "filename": p.name,
                "error": "Dosyadan metin cikarilamadi"}

    ftype = filetype_of(str(p))
    try:
        size = p.stat().st_size
    except OSError:
        size = 0
    preview = " ".join(text.split())[:220]

    with get_conn() as conn:
        cur = conn.execute(
            "INSERT INTO documents "
            "(filename, filetype, stored_path, size_bytes, char_count, "
            " chunk_count, summary) VALUES (?, ?, ?, ?, ?, ?, ?)",
            (p.name, ftype, stored_path or str(p), size, len(text),
             len(chunks), preview),
        )
        doc_id = cur.lastrowid
        for i, ch in enumerate(chunks):
            emb = embed_text(ch, config) if config else None
            emb_json = json.dumps(emb) if emb else None
            conn.execute(
                "INSERT INTO doc_chunks (doc_id, chunk_index, content, embedding) "
                "VALUES (?, ?, ?, ?)",
                (doc_id, i, ch, emb_json),
            )
    return {"ok": True, "id": doc_id, "filename": p.name, "filetype": ftype,
            "chunks": len(chunks), "chars": len(text), "preview": preview}


# ---------- Listeleme / silme ----------

def list_documents() -> list[dict]:
    """Yuklu belgeleri (en yeni once) listeler - arayuz paneli icin."""
    with get_conn() as conn:
        rows = conn.execute(
            "SELECT id, filename, filetype, size_bytes, char_count, "
            "chunk_count, summary, created_at FROM documents ORDER BY id DESC"
        ).fetchall()
    return rows_to_dicts(rows)


def all_document_ids() -> list[int]:
    with get_conn() as conn:
        rows = conn.execute("SELECT id FROM documents ORDER BY id DESC").fetchall()
    return [r["id"] for r in rows]


def delete_document(doc_id: int, remove_file: bool = True) -> None:
    """Bir belgeyi ve tum parcalarini siler. Diskteki kopyayi da (belgeler/
    klasorundeyse) siler."""
    with get_conn() as conn:
        row = conn.execute(
            "SELECT stored_path FROM documents WHERE id = ?", (doc_id,)
        ).fetchone()
        conn.execute("DELETE FROM doc_chunks WHERE doc_id = ?", (doc_id,))
        conn.execute("DELETE FROM documents WHERE id = ?", (doc_id,))
    if remove_file and row and row["stored_path"]:
        sp = row["stored_path"]
        # Sadece uygulamanin 'belgeler/' kopyasini sil; kullanicinin orijinal
        # dosyasina dokunma.
        if os.sep + "belgeler" + os.sep in sp or "/belgeler/" in sp:
            try:
                os.remove(sp)
            except OSError:
                pass


# ---------- Retrieval (RAG) ----------

def _all_chunks() -> list[dict]:
    with get_conn() as conn:
        rows = conn.execute(
            "SELECT c.id, c.doc_id, c.chunk_index, c.content, c.embedding, "
            "d.filename FROM doc_chunks c JOIN documents d ON d.id = c.doc_id "
            "ORDER BY c.doc_id DESC, c.chunk_index ASC"
        ).fetchall()
    return rows_to_dicts(rows)


def relevant_chunks(query: str | None, config: dict | None) -> list[dict]:
    """Soruyla anlamca en ilgili belge parcalarini secer (RAG).

    - Parca azsa: hepsini dondur.
    - Coksa: sorunun embedding'i ile benzerlige gore sirala, en iyi K'yi al.
    - Embedding yoksa: en yeni parcalari dondur (guvenli mod)."""
    chunks = _all_chunks()
    if not chunks:
        return []
    if len(chunks) <= SEMANTIC_THRESHOLD:
        return chunks[:TOP_CHUNKS]

    query_emb = embed_text(query, config) if (query and config) else None
    if not query_emb:
        return chunks[:TOP_CHUNKS]

    scored = []
    for c in chunks:
        if c.get("embedding"):
            try:
                emb = json.loads(c["embedding"])
                scored.append((cosine_similarity(query_emb, emb), c))
            except (json.JSONDecodeError, TypeError):
                continue
    if not scored:
        return chunks[:TOP_CHUNKS]
    scored.sort(key=lambda x: x[0], reverse=True)
    return [c for _, c in scored[:TOP_CHUNKS]]


def format_for_prompt(query: str | None = None,
                      config: dict | None = None) -> str:
    """Ilgili belge parcalarini LLM'in okuyacagi metne cevirir.
    Belge yoksa bos string doner (prompt sismesin)."""
    chunks = relevant_chunks(query, config)
    if not chunks:
        return ""
    lines, total = [], 0
    for c in chunks:
        piece = c["content"]
        if total + len(piece) > MAX_CONTEXT_CHARS:
            piece = piece[: max(0, MAX_CONTEXT_CHARS - total)]
        if not piece:
            break
        lines.append(f"[kaynak: {c['filename']}]\n{piece}")
        total += len(piece)
        if total >= MAX_CONTEXT_CHARS:
            break
    return "\n\n---\n\n".join(lines)
